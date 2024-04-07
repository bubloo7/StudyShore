from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def export_quiz(quiz_data, docx_filename):
    document = Document()

    heading_style = document.styles["Heading1"]
    heading_font = heading_style.font
    heading_font.size = Pt(14)  
    body_text_style = document.styles["BodyText"]
    body_text_font = body_text_style.font
    body_text_font.size = Pt(12)  

    document.add_heading("Quiz", level=1)

    for question_data in quiz_data:
        question = question_data["question"]
        possible_answers = question_data["possible_answers"]

        p = document.add_paragraph()
        p.add_run(question).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for answer in possible_answers:
            document.add_paragraph(answer, style="ListBullet")

    document.save(docx_filename)


# quiz_data = {
#     "quiz": [
#         {
#             "question": "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#             "possible_answers": [
#                 "Civil Rights Movement",
#                 "Photosynthesis",
#                 "Newton's Laws",
#                 "E=mc^2",
#             ],
#             "index": 1,
#         },
#         {
#             "question": "Mitochondria",
#             "possible_answers": [
#                 "Three fundamental principles of classical mechanics proposed by Sir Isaac Newton.",
#                 "Chemical formula for water, composed of two hydrogen atoms bonded to one oxygen atom.",
#                 "Organelles that generate most of the chemical energy needed to power the biochemical reactions of cells.",
#                 "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#             ],
#             "index": 2,
#         },
#         {
#             "question": "Three fundamental principles of classical mechanics proposed by Sir Isaac Newton.",
#             "possible_answers": [
#                 "Renaissance",
#                 "The Great Depression",
#                 "Civil Rights Movement",
#                 "Newton's Laws",
#             ],
#             "index": 3,
#         },
#         {
#             "question": "Chemical formula for water, composed of two hydrogen atoms bonded to one oxygen atom.",
#             "possible_answers": [
#                 "Cell Division",
#                 "The Great Depression",
#                 "Civil Rights Movement",
#                 "H2O",
#             ],
#             "index": 3,
#         },
#         {
#             "question": "A struggle for social justice that took place mainly during the 1950s and 1960s for African Americans to gain equal rights under the law in the United States.",
#             "possible_answers": [
#                 "Civil Rights Movement",
#                 "E=mc^2",
#                 "Palindrome",
#                 "The Great Depression",
#             ],
#             "index": 0,
#         },
#         {
#             "question": "Palindrome",
#             "possible_answers": [
#                 "Three fundamental principles of classical mechanics proposed by Sir Isaac Newton.",
#                 "The process by which a parent cell divides into two or more daughter cells.",
#                 "Organelles that generate most of the chemical energy needed to power the biochemical reactions of cells.",
#                 "A word, phrase, number, or other sequence of characters that reads the same forward and backward.",
#             ],
#             "index": 3,
#         },
#         {
#             "question": "The Great Depression",
#             "possible_answers": [
#                 "Three fundamental principles of classical mechanics proposed by Sir Isaac Newton.",
#                 "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#                 "A struggle for social justice that took place mainly during the 1950s and 1960s for African Americans to gain equal rights under the law in the United States.",
#                 "A severe worldwide economic depression that took place mostly during the 1930s.",
#             ],
#             "index": 3,
#         },
#         {
#             "question": "E=mc^2",
#             "possible_answers": [
#                 "Organelles that generate most of the chemical energy needed to power the biochemical reactions of cells.",
#                 "A severe worldwide economic depression that took place mostly during the 1930s.",
#                 "Albert Einstein's famous equation, which expresses the relationship between energy (E), mass (m), and the speed of light (c) squared.",
#                 "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#             ],
#             "index": 2,
#         },
#         {
#             "question": "Renaissance",
#             "possible_answers": [
#                 "The process by which a parent cell divides into two or more daughter cells.",
#                 "Albert Einstein's famous equation, which expresses the relationship between energy (E), mass (m), and the speed of light (c) squared.",
#                 "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#                 "A period in European history marking the transition from the Middle Ages to modernity and covering the 14th to 17th centuries.",
#             ],
#             "index": 3,
#         },
#         {
#             "question": "Cell Division",
#             "possible_answers": [
#                 "Chemical formula for water, composed of two hydrogen atoms bonded to one oxygen atom.",
#                 "The process by which a parent cell divides into two or more daughter cells.",
#                 "A severe worldwide economic depression that took place mostly during the 1930s.",
#                 "The process by which green plants and some other organisms use sunlight to synthesize foods with the help of chlorophyll.",
#             ],
#             "index": 1,
#         },
#     ]
# }["quiz"]
# export_quiz(quiz_data, "Quiz.docx")
