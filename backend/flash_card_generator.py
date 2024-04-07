from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def replace_text_in_table(doc, old_text, new_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell_format = cell.paragraphs[0].runs[0].font
                    cell_alignment = cell.paragraphs[0].alignment
                    cell.text = cell.text.replace(old_text, new_text)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = cell_format.size
                            run.font.bold = cell_format.bold
                            run.font.italic = cell_format.italic
                            run.font.underline = cell_format.underline
                        paragraph.alignment = cell_alignment


def duplicate_row(table, index, thing):
    original_row = table.rows[index]
    new_row = table.add_row()
    for i, cell in enumerate(original_row.cells):
        new_cell = new_row.cells[i]
        new_cell.text = cell.text[:-1] + str(thing)
        print(cell.text[:-1] + str(thing))
        new_cell.paragraphs[0].runs[0].font.size = cell.paragraphs[0].runs[0].font.size
        new_cell.paragraphs[0].runs[0].font.bold = cell.paragraphs[0].runs[0].font.bold
        new_cell.paragraphs[0].runs[0].font.italic = (
            cell.paragraphs[0].runs[0].font.italic
        )
        new_cell.paragraphs[0].runs[0].font.underline = (
            cell.paragraphs[0].runs[0].font.underline
        )

        new_cell.paragraphs[0].alignment = cell.paragraphs[0].alignment
        new_cell.vertical_alignment = cell.vertical_alignment


def export_flashcards(cards, docx_filename):

    doc = Document("stupid.docx")

    table = doc.tables[0]

    for i in range(len(cards) - 1):
        print("dupplicate")
        duplicate_row(table, i, i + 1)

    for i in range(len(cards)):
        replace_text_in_table(doc, "front" + str(i), cards[i][0])
        replace_text_in_table(doc, "back" + str(i), cards[i][1])

    for row in table.rows:
        row.height = Inches(1.7)

    for i in range(len(cards)):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w")))
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="F88A6B"/>'.format(nsdecls("w")))
        table.rows[i].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        table.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

    doc.save(docx_filename)


# notecards = [
#     ["Front of Card 1", "Back of Card 1"],
#     ["Front of Card 2", "Back of Card 2"],
#     ["Front of Card 3", "Back of Card 3"],
#     ["Front of Card 4", "Back of Card 4"],
#     ["Front of Card 5", "Back of Card 5"],
#     ["Front of Card 6", "Back of Card 6"],
#     ["Front of Card 7", "Back of Card 7"],
#     ["Front of Card 8", "Back of Card 8"],
# ]
# export_flashcards(notecards, "FlashCards.docx")
